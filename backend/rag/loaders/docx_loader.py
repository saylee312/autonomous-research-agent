from docx import Document


def load_docx(file_path):

    document = Document(file_path)

    result = {

        "text": [],

        "tables": [],

        "images": []
    }

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:

            result["text"].append(
                text
            )

    for table in document.tables:

        rows = []

        for row in table.rows:

            rows.append(
                [
                    cell.text
                    for cell in row.cells
                ]
            )

        result["tables"].append(
            rows
        )

    return result